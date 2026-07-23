# -*- coding: utf-8 -*-
"""Gera o relatório ANTECIPA (Word), documento endereçado à pesquisadora e à
orientadora. Requer python-docx (pip install python-docx) e as figuras geradas
por diagrama_calibracao.py e compara_explicacoes.py.

O .docx é gravado na RAIZ da pasta unb (um nível acima de antecipa-real), de
propósito FORA do repositório git: por ser documento nominal com avaliação
crítica, não é publicado (o conteúdo técnico impessoal está em RESULTADOS.md).
"""
from pathlib import Path

import docx
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(__file__).resolve().parent.parent          # .../unb/antecipa-real
PROC = BASE / "dados" / "processados"
DESTINO = BASE.parent / "Relatorio_ANTECIPA_dados_reais.docx"  # .../unb/ (fora do repo)

AZUL = RGBColor(0x1C, 0x5C, 0xAB)
CINZA = RGBColor(0x52, 0x51, 0x4E)
PRETO = RGBColor(0x0B, 0x0B, 0x0B)
FIGURA = str(PROC / "diagrama_calibracao.png")
FIGURA2 = str(PROC / "comparacao_explicacoes.png")

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)
for nome, tam, cor in [("Heading 1", 15, AZUL), ("Heading 2", 12, PRETO)]:
    h = doc.styles[nome]
    h.font.name = "Calibri"
    h.font.size = Pt(tam)
    h.font.bold = True
    h.font.color.rgb = cor

def p(texto="", bold=False, italic=False, cor=None, tam=None, center=False, just=True):
    par = doc.add_paragraph()
    if center:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif just:
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = par.add_run(texto)
    run.bold = bold
    run.italic = italic
    if cor: run.font.color.rgb = cor
    if tam: run.font.size = Pt(tam)
    return par

def rich(par, texto, bold=False, italic=False):
    run = par.add_run(texto)
    run.bold = bold
    run.italic = italic
    return run

def bullet(partes):
    par = doc.add_paragraph(style="List Bullet")
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if isinstance(partes, str):
        partes = [(partes, False)]
    for texto, bold in partes:
        r = par.add_run(texto)
        r.bold = bold
    return par

def hyperlink(par, texto, url):
    rel = par.part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK,
                             is_external=True)
    h = OxmlElement("w:hyperlink")
    h.set(qn("r:id"), rel)
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "1c5cab"); rpr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    r.append(rpr)
    t = OxmlElement("w:t"); t.text = texto; r.append(t)
    h.append(r)
    par._p.append(h)

def tabela(headers, rows, larguras_cm):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, htxt in enumerate(headers):
        cel = t.cell(0, j)
        cel.text = ""
        r = cel.paragraphs[0].add_run(htxt)
        r.bold = True
        r.font.size = Pt(9)
    for i, row in enumerate(rows, 1):
        for j, txt in enumerate(row):
            cel = t.cell(i, j)
            cel.text = ""
            r = cel.paragraphs[0].add_run(str(txt))
            r.font.size = Pt(9)
    for j, larg in enumerate(larguras_cm):
        for cel in t.columns[j].cells:
            cel.width = Cm(larg)
    doc.add_paragraph()
    return t

# ---------------------------------------------------------------- capa
p("RELATÓRIO TÉCNICO", bold=True, cor=CINZA, tam=11, center=True)
p("ANTECIPA — Da PoC simulada aos dados reais do PNCP e ao modelo preditivo",
  bold=True, cor=AZUL, tam=17, center=True)
p("Evolução do Produto Técnico-Tecnológico do projeto “Do Controle Reativo à Antecipação: "
  "IA para Gestão Preditiva de Riscos em Compras Públicas no STF”", italic=True, tam=11, center=True)
p("Para: Cristina Harumi Matsunaga e Profa. Dra. Marina Figueiredo Moreira (FACE/UnB)",
  cor=CINZA, tam=10, center=True)
p("Elaborado por David Nadler Prata, com apoio de IA (Claude) · 22 de julho de 2026",
  cor=CINZA, tam=10, center=True)

# ------------------------------------------------------- sumário executivo
doc.add_heading("Sumário executivo", level=1)
p("A Prova de Conceito do ANTECIPA, originalmente construída com dados simulados, foi evoluída em três "
  "saltos: (1) os três módulos do dashboard passaram a operar com dados reais e públicos (PNCP, "
  "compras.gov.br e dados abertos do CNPJ/Receita Federal); (2) foi coletado um dataset multi-órgão com "
  "25.488 contratos únicos de 11 órgãos do Judiciário federal, rotulado pelo desfecho real de cada "
  "contrato (aditivos, prorrogações e rescisões registrados nos termos do PNCP); e (3) foram treinados e "
  "avaliados modelos supervisionados em duas versões — a segunda com rótulo de horizonte fixo de 12 meses "
  "e probabilidades calibradas, hoje publicada no dashboard.")
p("Resultado central: o Gradient Boosting v2 atinge PR-AUC de 0,139 (6,6 vezes o acaso) e ROC-AUC de "
  "0,806 em validação temporal, capturando cerca de 4 vezes mais desfechos adversos que a automatização "
  "das regras da matriz de risco (baseline heurístico), com precisão comparável. Isso sustenta "
  "empiricamente a alegação central da dissertação — o sistema é preditivo, não um mero automatizador de "
  "regras — e responde diretamente às fragilidades 2.1, 2.2 e 2.7 da análise crítica. Todo o pipeline é "
  "reprodutível em ~5 minutos no Google Colab, sem credenciais, a partir do repositório público no GitHub.")
p("Achado complementar, e possivelmente o mais contundente para a defesa (seção 4.4): aplicada à carteira "
  "do STF apenas com dados anteriores à assinatura, a matriz da Res. 781/2022 apresenta precisão de 23,1% "
  "contra uma taxa base de 22,7% — isto é, não discrimina melhor que o acaso. No mesmo volume de alertas, "
  "o modelo supervisionado identifica 2,5 vezes mais contratos problemáticos. A automatização das regras, "
  "isoladamente, não entrega capacidade preditiva; é o aprendizado a partir dos desfechos reais que a "
  "produz.")

# ------------------------------------------------------- 1. o que foi construído
doc.add_heading("1. O que foi construído", level=1)
bullet([("Dashboard com dados reais", True),
        (" (ANTECIPA_dashboard_real.html): os três módulos do PTT — Spend Analysis, Raio-X do Fornecedor "
         "e Alertas de Risco — alimentados por APIs oficiais, preservando a identidade visual e os "
         "mecanismos de governança (human-in-the-loop) da PoC original. Arquivo autocontido: abre em "
         "qualquer navegador, sem servidor.", False)])
bullet([("Pipeline de coleta idempotente", True),
        (" em Python (somente biblioteca padrão): contratos e termos do STF e de 10 tribunais (STJ, TST, "
         "TSE, STM, CNJ, TRF2–TRF6), cadastros de 9.598 fornecedores na Receita Federal, preços de "
         "referência nacionais por item de catálogo — com cache em disco, retry/backoff e passe de reparo "
         "para janelas bloqueadas por rate-limit.", False)])
bullet([("Dois datasets rotulados pelo desfecho real", True),
        (": v1 (25.488 contratos, rótulo sem horizonte) e v2 (16.783 contratos elegíveis, rótulo de "
         "evento adverso em até 12 meses da assinatura — correção da censura à direita).", False)])
bullet([("Modelos supervisionados com validação temporal", True),
        (" (Regressão Logística e Gradient Boosting), comparados a um baseline heurístico transparente "
         "que automatiza as escalas da Res. 781/2022; experimentos adicionais confrontando treino só-STF "
         "contra multi-órgão (com bootstrap) e duas técnicas de explicação (SHAP sobre o modelo em uso "
         "contra modelo substituto linear).", False)])
bullet([("Explicabilidade auditada por evidência", True),
        (": as explicações do painel usam SHAP sobre o próprio Gradient Boosting, escolha adotada após "
         "medir que o modelo substituto anteriormente empregado coincidia com o modelo real quanto ao "
         "fator principal em apenas 12,3% dos contratos (seção 4.5).", False)])
bullet([("Publicação e reprodutibilidade", True),
        (": repositório público no GitHub com código e caches de dados, dois notebooks para Google Colab "
         "(um monta o Google Drive; outro clona o repositório e roda sem credenciais) e script de "
         "sincronização Drive→GitHub.", False)])

# ------------------------------------------------------- 2. dados
doc.add_heading("2. Dados coletados (todos reais e públicos)", level=1)
tabela(
    ["Fonte", "Conteúdo", "Volume"],
    [
        ["API de Consulta do PNCP", "Contratos do STF e de 10 tribunais federais desde 2023 (objeto, valores, vigência, fornecedor)", "31.907 registros; 25.488 contratos únicos"],
        ["API do PNCP — termos", "Aditivos, apostilamentos e rescisões de cada contrato, com data e valores", "Milhares de termos; base do rótulo"],
        ["API de busca do PNCP", "Presença nacional (todos os órgãos) dos maiores fornecedores do STF", "Top-15 da carteira STF"],
        ["Dados abertos do CNPJ (Receita Federal, via minhareceita.org)", "Capital social, porte, idade, CNAE, situação cadastral, natureza jurídica", "9.598 de 9.600 fornecedores únicos"],
        ["Pesquisa de preços do compras.gov.br + catálogo de PDMs", "Distribuição nacional de preços homologados por item (Módulo 1)", "7 itens com amostra nacional"],
    ],
    [4.5, 8.0, 4.0])
p("Recorte temporal: vigência plena da Lei 14.133/2021 (2023 em diante), conforme decisão metodológica "
  "do projeto. No STF: 365 registros de contrato (R$ 1,195 bilhão em valor global), 237 termos aditivos e "
  "22 rescisões/extinções — 22,7% dos contratos com algum desfecho adverso observável.")

# ------------------------------------------------------- 3. metodologia
doc.add_heading("3. Rotulagem pelo desfecho real e modelagem", level=1)
doc.add_heading("3.1 Rótulo (resposta à fragilidade 2.1 — circularidade)", level=2)
p("Em vez de rotular pelos escores da própria matriz de risco (o que faria o modelo apenas reproduzir as "
  "regras), cada contrato é rotulado pelo que de fato aconteceu depois da assinatura, extraído dos termos "
  "publicados no PNCP: rescisão/extinção registrada, acréscimo de valor superior a 25% ou prorrogação "
  "superior a 365 dias.")
p("Na v2, o rótulo ganhou horizonte fixo: evento adverso em até 12 meses da assinatura (datado pelo "
  "termo), e só entram no dataset contratos com 12 meses completos de observação. Essa correção eliminou "
  "um artefato importante: na v1, a prevalência “caía” de 5,2% no treino para 1,6% no teste não por "
  "melhora real, mas porque contratos recentes ainda não tiveram tempo de acumular termos (censura à "
  "direita). Com o horizonte fixo, a prevalência fica estável por coorte de assinatura: 1,5% (2023), "
  "1,7% (2024) e 2,1% (2025).")
doc.add_heading("3.2 Features (somente informação disponível na assinatura)", level=2)
p("Valor global, duração prevista, categoria do processo, órgão, capital social, idade do CNPJ, porte, "
  "opção pelo Simples, situação cadastral, natureza de consórcio, razão capital/valor e recorrência do "
  "fornecedor. Nenhuma variável posterior à assinatura entra como feature — o desfecho é usado apenas "
  "como rótulo.")
doc.add_heading("3.3 Protocolo de avaliação (resposta à fragilidade 2.7)", level=2)
p("Validação temporal (treino 2023–2024, teste 2025; contratos de 2026 excluídos da avaliação na v1 por "
  "rótulo imaturo), métricas para classe rara — precision, recall e F1 da classe positiva, PR-AUC e "
  "ROC-AUC; nunca acurácia. Limiar de decisão escolhido no treino (máximo F1 em validação cruzada), "
  "jamais no teste. Baseline heurístico avaliado nas mesmas condições.")

doc.add_heading("3.4 As três leituras de risco e como cada uma é calculada", level=2)
p("O painel apresenta cada contrato sob três leituras distintas, que respondem a perguntas diferentes e "
  "não devem ser confundidas. Duas delas são ex-ante (usam apenas o que se sabia no dia da assinatura) e "
  "uma é retrospectiva. A distinção é essencial para interpretar corretamente os resultados da seção 4.")

p("(a) Nível ex-ante — Res. 781/2022. ", bold=True)
p("Escore heurístico transparente que parte de 1,0 ponto de probabilidade e soma apenas fatores "
  "conhecidos na assinatura:")
bullet("capital social inferior a 2% do volume contratado nacional: +1,2; superior a 30%: −0,5; "
       "não informado na base aberta (comum em consórcios): +0,3")
bullet("CNPJ com menos de 3 anos: +0,8; com mais de 10 anos: −0,4")
bullet("crescimento contratual atípico (3 ou mais vezes a mediana histórica no PNCP): +1,0")
bullet("situação cadastral na Receita diferente de “Ativa”: +1,5")
bullet("vigência prevista superior a 12 meses (serviço continuado): +0,4")
p("O total é arredondado e limitado ao intervalo 1–5 da escala normativa.")

p("(b) Probabilidade prevista na assinatura — modelo supervisionado. ", bold=True)
p("Gradient Boosting com calibração isotônica (5 folds), de modo que a porcentagem exibida corresponde a "
  "frequência observada, e não a um escore arbitrário:")
bullet("treino: 16.783 contratos de 11 órgãos do Judiciário federal (2022–2025), todos com 12 meses "
       "completos de observação")
bullet("rótulo: desfecho adverso em até 12 meses da assinatura — rescisão/extinção, acréscimo de valor "
       "superior a 25% ou prorrogação superior a 365 dias")
bullet("variáveis: apenas as disponíveis na assinatura — valor global, duração, categoria, órgão, capital "
       "social, idade e porte do fornecedor, situação cadastral, razão capital/valor e recorrência no PNCP")
bullet("na matriz de riscos, a probabilidade contínua é convertida para a escala 1–5 por faixas fixas "
       "(1: abaixo de 2%; 2: de 2 a 5%; 3: de 5 a 10%; 4: de 10 a 20%; 5: 20% ou mais), ancoradas na taxa "
       "base e não em quantis, para permanecerem comparáveis entre períodos e órgãos")
p("A explicação por contrato utiliza valores SHAP (TreeExplainer) calculados sobre o próprio Gradient "
  "Boosting — o modelo que gera a probabilidade —, em escala de log-odds. As contribuições de variáveis "
  "categóricas são somadas por variável original, de modo que não sejam exibidas atribuições a níveis "
  "que o contrato não assume. A seção 4.5 documenta a evidência que motivou essa escolha.")

p("(c) Nível observado — Res. 781/2022. ", bold=True)
p("O mesmo escore da leitura (a), acrescido do que já ocorreu com o contrato após a assinatura:")
bullet("dois ou mais termos aditivos firmados: +0,8")
bullet("valor acumulado mais de 10% acima do inicial: +0,8")
bullet("termo de rescisão/extinção registrado: +1,5")
p("Trata-se, portanto, de leitura retrospectiva: parte do que ela mede já é o próprio problema "
  "materializado, e não uma previsão. É a razão pela qual um contrato pode figurar como “Crítico” nesta "
  "leitura e apresentar probabilidade modesta nas leituras ex-ante.")

p("Elementos comuns e ressalva. ", bold=True)
p("Nas leituras (a) e (c), o eixo de impacto (1–5) é o quintil do valor global do contrato dentro da "
  "carteira do STF, e o nível resulta do produto probabilidade × impacto nas faixas da norma: até 4 "
  "Baixo; até 9 Moderado; até 16 Elevado; acima de 16 Crítico. As três leituras compartilham o mesmo eixo "
  "de impacto, o que as torna diretamente comparáveis. Registre-se, contudo, que o escore ex-ante é uma "
  "reconstrução realizada no presente: utiliza os dados cadastrais atuais da Receita Federal, que podem "
  "ter se alterado desde a data de assinatura. A versão rigorosa exigiria séries históricas de dados "
  "cadastrais, não disponíveis na base aberta.")

# ------------------------------------------------------- 4. resultados
doc.add_heading("4. Resultados", level=1)
doc.add_heading("4.1 Modelo principal (teste temporal 2025)", level=2)
tabela(
    ["Modelo", "Precision", "Recall", "F1", "PR-AUC", "ROC-AUC"],
    [
        ["Gradient Boosting v2 (12 meses)", "0,109", "0,494", "0,179", "0,139 (lift 6,6×)", "0,806"],
        ["Regressão Logística v2", "0,052", "0,647", "0,096", "0,065", "0,755"],
        ["Gradient Boosting v1", "0,064", "0,393", "0,111", "0,071 (lift 4,4×)", "0,798"],
        ["Baseline heurístico (Res. 781)", "0,067", "0,129", "0,088", "0,033 (lift 2×)", "0,705"],
    ],
    [5.2, 2.2, 2.0, 1.8, 3.2, 2.1])
p("Leitura: o modelo aprende sinal real que as regras não capturam — com precisão comparável à do "
  "baseline, encontra quase 4 vezes mais desfechos adversos. A v2 domina a v1 em todas as métricas "
  "relevantes (a comparação correta entre versões é pelo lift sobre o acaso, pois testes e rótulos "
  "diferem). Em visão operacional de fila de diligência: entre os 100 contratos de 2025 com maior escore "
  "v2, 16 tiveram desfecho adverso real em 12 meses — 7,6 vezes a taxa base — capturando 19% de todos os "
  "adversos do ano.")
doc.add_heading("4.2 Experimento: treinar só com STF × multi-órgão (fragilidade 2.2)", level=2)
p("Ambos avaliados no mesmo teste (contratos do STF assinados em 2025: 106 contratos, 13 adversos). O "
  "modelo só-STF (148 contratos de treino) obteve PR-AUC 0,384; o multi-órgão (12.483 contratos), 0,319. "
  "O bootstrap (4.000 reamostras) mostra que a diferença não é conclusiva: Δ = +0,069 com IC95% "
  "[−0,221; +0,350]. Interpretação equilibrada: o padrão local do STF carrega sinal específico relevante "
  "(sua carteira tem prevalência de desfecho adverso muito superior à dos demais órgãos), mas o modelo "
  "multi-órgão entrega desempenho comparável com 84 vezes mais dados e muito mais estabilidade. "
  "Recomendação com nuance empírica própria: treinar no padrão nacional e calibrar/afinar por órgão.")
doc.add_heading("4.3 Calibração (o que o dashboard exibe)", level=2)
p("O treino com pesos de classe (class_weight='balanced') instrui o modelo a tratar as classes como se "
  "fossem equilibradas, quando na realidade apenas 1,8% dos contratos têm desfecho adverso. Isso melhora "
  "o aprendizado em classe rara, mas distorce a saída: o número produzido é um escore inflado, não uma "
  "probabilidade. Para corrigir, aplica-se calibração isotônica — uma regressão monotônica não "
  "decrescente que aprende, a partir dos próprios dados, o mapeamento entre escore e frequência "
  "efetivamente observada, sem impor formato de curva. O ajuste é feito por validação cruzada em 5 "
  "partições, de modo que a calibração não utilize as mesmas observações do treino.")
p("A tabela e a figura abaixo, medidas no conjunto de teste temporal (contratos assinados em 2025, nunca "
  "vistos no treino), demonstram o efeito:")
tabela(
    ["Métrica", "Sem calibração", "Calibrado (isotônica)", "Referência"],
    [
        ["Brier score (menor é melhor)", "0,1280", "0,0196", "—"],
        ["ECE — erro de calibração esperado", "0,2240", "0,0063", "0 (ideal)"],
        ["Probabilidade média prevista", "24,5%", "2,0%", "2,1% (frequência real)"],
        ["ROC-AUC (discriminação)", "0,8064", "0,8019", "—"],
    ],
    [5.6, 3.2, 3.9, 3.8])
p("Três leituras desta tabela merecem destaque. Primeiro, o modelo não calibrado prevê em média 24,5% "
  "quando a frequência real é 2,1% — superestimativa de aproximadamente doze vezes, que produziria "
  "alarmes falsos sistemáticos se exibida ao servidor. Segundo, após a calibração a média passa a 2,0%, "
  "praticamente idêntica à frequência observada, e o erro de calibração (ECE) cai de 0,224 para 0,006 — "
  "redução de trinta e seis vezes. Terceiro, e decisivo para a validade do procedimento, o ROC-AUC "
  "permanece essencialmente inalterado (0,8064 contra 0,8019): como a transformação é monotônica, a "
  "ordenação dos contratos é preservada. A calibração não melhora nem piora a capacidade de discriminar; "
  "ela apenas torna a escala interpretável — que é precisamente o que se pretende ao exibir uma "
  "porcentagem a quem decide.")
try:
    doc.add_picture(str(FIGURA), width=Cm(16.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p("Figura 1 — Diagrama de confiabilidade no teste temporal de 2025. Em (a), escala completa: a curva "
      "sem calibração (vermelha) afasta-se drasticamente da diagonal, chegando a prever 80% onde a "
      "frequência real é 10%. Em (b), zoom na faixa efetivamente ocupada pelo modelo calibrado (azul), "
      "que acompanha a diagonal de perto. Em (c), a distribuição das probabilidades evidencia o "
      "deslocamento: o modelo bruto espalha contratos por toda a escala, enquanto o calibrado os "
      "concentra próximo à prevalência real.", italic=True, cor=CINZA, tam=9)
except Exception as e:  # figura ainda não gerada
    p(f"[Figura do diagrama de confiabilidade não incorporada: {e}]", italic=True, cor=CINZA, tam=9)
p("As explicações por contrato usam valores SHAP calculados sobre o próprio Gradient Boosting, "
  "apresentados explicitamente como insumo para a motivação humana, e não como a motivação em si "
  "(fragilidade 2.8). A seção seguinte documenta a evidência que levou a substituir a abordagem "
  "anteriormente adotada.")

doc.add_heading("4.4 Aplicação à carteira do STF: a matriz ex-ante não discrimina", level=2)
p("Aplicando os dois critérios de alerta aos 365 contratos do STF publicados no PNCP (taxa base de "
  "desfecho adverso: 22,7%), obtém-se o resultado abaixo. O escore ex-ante é o mesmo escore heurístico "
  "da Res. 781/2022, porém calculado apenas com os fatores conhecidos no dia da assinatura — a leitura "
  "que o sistema teria exibido ao gestor antes de contratar.")
tabela(
    ["Critério de alerta", "Alertas", "Acertos", "Precisão", "Recall", "Lift"],
    [
        ["Ex-ante (Res. 781, Elevado/Crítico)", "52", "12", "23,1%", "14,5%", "1,01×"],
        ["ML ≥ 20%", "100", "57", "57,0%", "68,7%", "2,51×"],
        ["ML ≥ 10%", "159", "65", "40,9%", "78,3%", "1,80×"],
        ["ML ≥ 5%", "232", "74", "31,9%", "89,2%", "1,40×"],
        ["ML top-52 (mesmo volume do ex-ante)", "52", "30", "57,7%", "36,1%", "2,54×"],
        ["(referência) Observado, Elevado/Crítico", "66", "22", "33,3%", "26,5%", "1,47×"],
    ],
    [5.8, 1.9, 1.9, 1.9, 1.8, 1.7])
p("O achado central: a heurística ex-ante apresenta precisão de 23,1% contra uma taxa base de 22,7% — "
  "lift de 1,01×. Nesta carteira ela não discrimina: sortear 52 contratos ao acaso acertaria "
  "praticamente o mesmo (esperado ~11,8; obtido 12), diferença inteiramente dentro do ruído estatístico. "
  "Comparados no mesmo custo operacional (52 alertas para cada método), o modelo identifica 30 contratos "
  "problemáticos contra 12 da régua normativa — 2,5 vezes mais.")
p("Três ressalvas indispensáveis para a leitura correta desta tabela:", bold=True)
bullet([("Viés in-sample nos números do ML. ", True),
        ("Os contratos do STF integraram o treino do modelo calibrado (2022–2025, 11 órgãos), de modo "
         "que os valores acima são otimistas. A medição fora da amostra é a da seção 4.1: precisão de "
         "10,9% e recall de 49,4% no teste temporal de 2025 — nominalmente menor, mas sobre prevalência "
         "de 2,1%, o que corresponde a lift de aproximadamente 5×.", False)])
bullet([("Rótulo mais frouxo. ", True),
        ("Esta tabela usa o rótulo do dashboard (“teve desfecho adverso a qualquer momento”), não o de "
         "horizonte de 12 meses da avaliação formal. Rótulo mais comum eleva a precisão nominal e "
         "dificulta a discriminação fina — por isso a coluna comparável entre contextos é o lift, não a "
         "precisão absoluta.", False)])
bullet([("Por que a heurística falha justamente no STF. ", True),
        ("A carteira do Tribunal é homogeneamente de alto risco pelos critérios da norma: quase tudo é "
         "serviço continuado, longo e de valor elevado. As variáveis que a heurística usa para "
         "discriminar (quantil de valor, vigência superior a 12 meses) estão saturadas — se todos "
         "pontuam, ninguém se destaca. Numa base multi-órgão heterogênea ela ainda apresentava algum "
         "sinal (lift próximo de 4× na avaliação formal); dentro do STF, esse sinal desaparece.", False)])
p("Este é, na avaliação técnica deste relatório, o resultado mais forte do trabalho para a defesa: "
  "automatizar a matriz da Res. 781/2022 com dados anteriores à assinatura não produz capacidade "
  "preditiva na carteira do STF — e é precisamente essa lacuna que o modelo supervisionado preenche. "
  "Reforça-se, porém, a contrapartida operacional: 100 a 160 alertas numa carteira de 365 contratos "
  "excedem a capacidade de diligência de qualquer equipe, o que sustenta o uso do modelo como fila de "
  "priorização (top-N por probabilidade) em vez de alerta binário.")

doc.add_heading("4.5 Escolha da técnica de explicação: SHAP × modelo substituto", level=2)
p("A primeira versão do painel exibia, como explicação de cada contrato, as contribuições aditivas de uma "
  "regressão logística treinada sobre os mesmos dados. A escolha era pragmática — a regressão logística é "
  "aditiva por construção, o que produz explicação exata e auditável sem bibliotecas adicionais —, mas "
  "envolvia uma fragilidade conceitual: a logística é um modelo substituto, que explica a si mesmo e não "
  "o Gradient Boosting do qual provém a probabilidade exibida.")
p("Para decidir com evidência, e não por preferência teórica, as duas formas de explicação foram "
  "calculadas para os mesmos contratos do STF e no mesmo espaço de variáveis (ambos os modelos ajustados "
  "sobre a matriz pré-processada idêntica). Os resultados são desfavoráveis ao modelo substituto:")
tabela(
    ["Indicador de concordância entre as duas explicações", "Valor"],
    [
        ["Correlação de Spearman média entre as atribuições", "0,222"],
        ["Contratos em que o fator principal coincide", "12,3%"],
        ["Sobreposição média entre os três principais fatores", "56,3%"],
        ["Contratos analisados", "293"],
    ],
    [10.5, 4.5])
p("Em outras palavras, em aproximadamente nove de cada dez contratos o fator apresentado como principal "
  "não era o principal do modelo que produziu a porcentagem. A inspeção das atribuições revelou a causa "
  "estrutural: em variáveis codificadas por indicadores binários (one-hot), a contribuição da regressão "
  "logística é o coeficiente multiplicado por um — isto é, uma constante idêntica para todos os contratos "
  "do mesmo órgão. O indicador “órgão: STF” aparecia assim como fator principal em praticamente todos os "
  "casos, com magnitude fixa, dominando a explicação sem discriminar coisa alguma entre contratos. A "
  "importância global confirma a inversão: a variável de maior peso no Gradient Boosting (duração do "
  "contrato, com contribuição média absoluta de 1,521) tinha peso quase dez vezes menor na regressão "
  "logística (0,156), enquanto “órgão: STF” liderava nesta última (3,024) contra 0,810 no modelo real.")
try:
    doc.add_picture(str(FIGURA2), width=Cm(16.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p("Figura 2 — A mesma predição explicada de duas formas, para três contratos do STF. À esquerda, SHAP "
      "sobre o Gradient Boosting; à direita, contribuições da regressão logística. Note-se que “órgão: "
      "STF” figura como fator principal da logística nos três casos, com magnitude idêntica (+3,02), "
      "inclusive no contrato de menor risco — onde o SHAP indica que o fator dominante é o número de "
      "contratos do fornecedor, reduzindo o risco. A explicação substituta chegava a contradizer a "
      "direção do raciocínio do modelo.", italic=True, cor=CINZA, tam=9)
except Exception as e:
    p(f"[Figura da comparação de explicações não incorporada: {e}]", italic=True, cor=CINZA, tam=9)
p("Diante desse resultado, as explicações do painel foram migradas para valores SHAP (TreeExplainer) "
  "calculados sobre o próprio Gradient Boosting, atendendo ao que o projeto de qualificação já previa. "
  "Duas providências acompanharam a migração. Primeiro, as contribuições dos níveis de cada variável "
  "categórica passaram a ser somadas por variável original: sem isso, o SHAP atribuía valor a níveis que "
  "o contrato não assume (por exemplo, “órgão: TSE” em contrato do STF, medindo a contribuição de não ser "
  "do TSE), o que é tecnicamente correto mas incompreensível para o usuário. Segundo, mede-se e reporta-se "
  "a fidelidade da explicação ao modelo publicado.")
p("Ressalva de fidelidade. ", bold=True)
p("O SHAP explica o Gradient Boosting ajustado sobre toda a base de treino, ao passo que a probabilidade "
  "exibida provém do mesmo algoritmo calibrado em cinco partições. A concordância de ordenação entre os "
  "dois é de 0,878 (Spearman) — muito superior aos 0,222 do modelo substituto, porém inferior à unidade. "
  "Parte dessa diferença é artefato de empates, e não discordância efetiva: a calibração isotônica é "
  "função em degraus e reduz os 293 contratos do STF a apenas 111 valores distintos de probabilidade, o "
  "que penaliza o coeficiente de Spearman sem que haja divergência real de ordenação.")
p("O indicador é recalculado a cada publicação do modelo e registrado no arquivo de predições para fins "
  "de auditoria. Além disso, o próprio painel exibe a ressalva no bloco “Como as três leituras são "
  "calculadas”, de modo que quem consulta o dashboard conhece o limite da explicação sem depender desta "
  "documentação — coerente com a orientação de que a explicabilidade deve estar disponível a quem decide, "
  "e não apenas ao corpo técnico.")
p("Eliminar por completo o resíduo seria possível calibrando com conjunto reservado (parâmetro "
  "cv=\"prefit\"): haveria então um único modelo, e o SHAP explicaria exatamente o escore que alimenta a "
  "calibração. O custo seria retirar observações do ajuste do modelo, num conjunto que dispõe de apenas "
  "296 casos positivos. A configuração adotada privilegia a qualidade da calibração (ECE de 0,006, "
  "seção 4.3) e assume um resíduo de fidelidade que é medido e divulgado, em vez de ignorado. Trata-se "
  "de escolha entre duas propriedades desejáveis, com o critério explicitado.")
p("Registre-se, por fim, que a alternativa oposta permanece defensável e merece discussão na dissertação: "
  "publicar a própria regressão logística como modelo, abrindo mão de aproximadamente metade do poder "
  "preditivo (PR-AUC de 0,065 contra 0,139) em troca de um sistema em que a explicação é o modelo, sem "
  "substituto nem aproximação. Em contexto de Direito Administrativo, no qual o princípio da motivação "
  "pode pesar mais que pontos de acurácia, essa troca é argumentável. O que não se sustenta é a "
  "combinação anteriormente adotada — modelo de árvores explicado por modelo linear distinto.")

# ------------------------------------------------------- 5. achados
doc.add_heading("5. Achados de engenharia de dados (material para a dissertação)", level=1)
bullet([("O STF não preenche o código CATMAT/CATSER", True),
        (" nos itens que publica no PNCP: 0 de 916 itens tinham o campo. A comparação nacional de preços "
         "do Módulo 1 exigiu reconstruir a ponte descrição → PDM → códigos de item via catálogo do "
         "compras.gov.br. Confirma na prática o item 2.4 da análise crítica: a maior parte do esforço é "
         "engenharia e limpeza de dados.", False)])
bullet([("As listas paginadas do PNCP contêm duplicatas", True),
        (" (republicações/retificações): o TST retorna 13.765 registros para 9.407 contratos únicos. "
         "Deduplicação por numeroControlePncp é obrigatória em qualquer uso analítico do PNCP.", False)])
bullet([("CNPJs raiz agregam unidades", True),
        (": TST e TSE, pelo CNPJ raiz, trazem as unidades regionais da Justiça do Trabalho e Eleitoral — "
         "útil para volume, mas exige atenção ao interpretar “órgão”.", False)])
bullet([("Serviços continuados são publicados como item único de valor global", True),
        (" (quantidade 1), o que inviabiliza comparação unitária de preços de serviços sem normalização "
         "adicional (postos, m² etc.).", False)])
bullet([("Consórcios aparecem com capital social zero", True),
        (" na base aberta da Receita — é dado não informado, não “empresa sem lastro”; o escore trata o "
         "caso separadamente para não gerar falso alerta.", False)])
bullet([("A API do PNCP aplica rate-limit", True),
        (" (HTTP 429) em coletas longas; o pipeline usa retry com backoff e um passe de reparo. Pequenas "
         "lacunas residuais estão documentadas no repositório.", False)])

# ------------------------------------------------------- 6. análise crítica
doc.add_heading("6. Correspondência com a análise crítica", level=1)
tabela(
    ["Recomendação", "Status"],
    [
        ["1. Rotular pelo desfecho real, não pela matriz", "Implementada (v1 e v2; horizonte de 12 meses na v2)"],
        ["2. Treinar multi-órgão e aplicar ao STF", "Implementada + experimento próprio com bootstrap (seção 4.2)"],
        ["3. APIs oficiais em vez de raspagem", "Implementada; sanções (CEIS/CNEP, SICAF, CNDT) pendentes de chave institucional"],
        ["4. Dimensionar limpeza/normalização CATMAT", "Confirmada na prática; ponte descrição→PDM construída para 7 itens"],
        ["5. Normativo: alerta dispara diligência, nunca inabilitação", "Institucional — reforçada nos textos do dashboard e relatórios"],
        ["6. Monitorar concordância humano-máquina", "Pendente (exige uso real com servidores)"],
        ["7. Precision/recall/F1 + calibração em vez de acurácia", "Implementada (validação temporal, PR-AUC, calibração isotônica)"],
        ["8. SHAP como insumo da motivação, não a motivação", "Implementada: SHAP sobre o modelo publicado (seção 4.5), apresentado como insumo"],
        ["9. Explicitar arquitetura batch vs. tempo real", "Explicitada: pipeline batch regenerável; dashboard autocontido"],
        ["10. Dono institucional e sustentação", "Pendente (decisão institucional)"],
    ],
    [8.2, 8.3])

# ------------------------------------------------------- 7. nota 781
doc.add_heading("7. Nota importante: operacionalização da Res. 781/2022", level=1)
p("A Res. 781/2022 fornece a moldura (escalas 1–5 de probabilidade e impacto, classificação pelo "
  "produto), mas não fornece uma tabela que converta dados observáveis em pontos. A tabela de pesos usada "
  "no baseline heurístico (ex.: capital/volume < 2% → +1,2; empresa < 3 anos → +0,8) é uma "
  "operacionalização proposta neste trabalho — plausível, porém arbitrária. Recomenda-se fortemente: "
  "(a) obter internamente o texto e anexos da resolução; (b) validar/calibrar o quadro de mapeamento "
  "(descritor da norma → proxy de dados → pontos) com os gestores de risco do STF, no processo de "
  "co-design já previsto no projeto; (c) documentar esse quadro na dissertação. Importante: o resultado "
  "central (modelo supervisionado) não depende dessa tabela — o rótulo vem do desfecho real, justamente "
  "para escapar da circularidade.")

# ------------------------------------------------------- 8. limitações
doc.add_heading("8. Limitações honestas", level=1)
bullet("O rótulo captura desfechos administráveis (aditivos, prorrogações, rescisões) — não fraude ou "
       "conluio, raros e não observáveis nestas bases.")
bullet("O horizonte de 12 meses não captura desfechos tardios (ex.: aditivos do 2º ano em serviços "
       "continuados); análise de sensibilidade com 18/24 meses é extensão natural.")
bullet("A feature de recorrência do fornecedor é contada no dataset completo (leve vazamento temporal; "
       "em produção, contar apenas contratos anteriores à assinatura).")
bullet("A precisão absoluta dos alertas é baixa (classe rara) — o uso correto é ranqueamento com fila "
       "por capacidade de diligência, e a decisão é sempre humana.")
bullet("Comparações de preço por PDM agregam especificações heterogêneas: desvios disparam diligência, "
       "nunca conclusão de sobrepreço.")
bullet("As explicações SHAP descrevem o Gradient Boosting ajustado sobre toda a base, enquanto a "
       "probabilidade exibida provém do mesmo algoritmo calibrado em cinco partições — concordância de "
       "ordenação de 0,878 (Spearman), parcialmente reduzida por empates da calibração (seção 4.5).")

# ------------------------------------------------------- 9. próximos passos
doc.add_heading("9. Próximos passos recomendados (em ordem de custo-benefício)", level=1)
bullet([("Chave gratuita da API do Portal da Transparência", True),
        (" (cadastro por e-mail) → CEIS/CNEP reais no Módulo 2 e possível enriquecimento do rótulo com "
         "sanções aplicadas.", False)])
bullet([("Validar a operacionalização da Res. 781 com os gestores", True),
        (" (seção 7) e registrar o quadro de mapeamento na dissertação.", False)])
bullet([("Custo assimétrico explícito", True),
        (" (falso negativo em contrato crítico vs. falso positivo) como peso amostral — transforma o "
         "balanceamento em decisão de gestão documentável (item 2.7).", False)])
bullet([("Sensibilidade do horizonte", True),
        (" (18/24 meses) e fine-tuning por órgão sobre o modelo nacional (seção 4.2).", False)])
bullet([("Piloto com servidores", True),
        (": medir a taxa de concordância humano-máquina (recomendação nº 6) e ajustar o desenho "
         "comportamental dos botões de governança.", False)])

# ------------------------------------------------------- 10. recursos
doc.add_heading("10. Recursos e reprodutibilidade", level=1)
par = doc.add_paragraph(); rich(par, "Repositório público (código + dados + relatórios): ", bold=True)
hyperlink(par, "github.com/DavidNadlerPrata/antecipa-pncp",
          "https://github.com/DavidNadlerPrata/antecipa-pncp")
par = doc.add_paragraph(); rich(par, "Reprodução em ~5 min no Google Colab (sem credenciais): ", bold=True)
hyperlink(par, "abrir ANTECIPA_colab_github.ipynb no Colab",
          "https://colab.research.google.com/github/DavidNadlerPrata/antecipa-pncp/blob/main/ANTECIPA_colab_github.ipynb")
par = doc.add_paragraph(); rich(par, "Dashboard publicado (link direto): ", bold=True)
hyperlink(par, "davidnadlerprata.github.io/antecipa-pncp/antecipa-4a8g1jdey1/",
          "https://davidnadlerprata.github.io/antecipa-pncp/antecipa-4a8g1jdey1/")
p("Endereço não indexado por buscadores (meta noindex) e de URL não adivinhável — adequado para "
  "compartilhamento com a orientadora e a banca antes da defesa. Observe-se que a proteção é por "
  "obscuridade, não por autenticação: quem tiver o link acessa. O arquivo também está no Drive em "
  "unb/ANTECIPA_dashboard_real.html (autocontido, abre em qualquer navegador).",
  italic=True, cor=CINZA, tam=9.5)
p("No Módulo 3, o dashboard apresenta as três leituras lado a lado — nível ex-ante e nível observado "
  "(ambos pela Res. 781/2022) e probabilidade prevista pelo modelo —, com as três matrizes de risco "
  "correspondentes e o detalhamento por contrato, que separa visualmente os fatores conhecidos na "
  "assinatura dos posteriores a ela.")
par = doc.add_paragraph(); rich(par, "Relatórios detalhados no repositório: ", bold=True)
rich(par, "README.md (visão geral e achados), dados/processados/relatorio_modelo.md (v1), "
          "relatorio_modelo_v2.md (v2) e comparacao_stf_vs_multi.md (experimento com bootstrap).")
p("Todos os dados utilizados são públicos (PNCP, compras.gov.br, dados abertos do CNPJ). Nenhuma "
  "informação interna do STF foi acessada; nada aqui constitui parecer jurídico. O repositório pode ser "
  "tornado privado ou transferido para a conta da pesquisadora a qualquer momento.",
  italic=True, cor=CINZA, tam=9.5)

doc.save(str(DESTINO))
print(f"gerado: {DESTINO} ({DESTINO.stat().st_size/1024:.0f} KB)")
